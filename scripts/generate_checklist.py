# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page setup (A4) ---
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.2)
section.right_margin  = Cm(2.2)
section.top_margin    = Cm(2.2)
section.bottom_margin = Cm(2.2)

# --- Color palette ---
NAVY    = (15, 40, 90)
GOLD    = (184, 134, 11)
DARK    = (30, 30, 30)
GREY    = (100, 100, 100)
WHITE   = (255, 255, 255)
LIGHT_BG = (245, 248, 255)

# --- Helpers ---
def make_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def rtl_run(para, text, bold=False, size=11, color=DARK, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rtl_elem = OxmlElement('w:rtl')
    rPr.append(rtl_elem)
    return run

def shade_para(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def set_para_border_bottom(paragraph, color='C0C0C0', sz='4'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_heading(doc, number, title_he, title_en):
    """Big section header with colored background"""
    p = doc.add_paragraph()
    make_rtl(p)
    shade_para(p, '0F285A')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    rtl_run(p, f'פרק {number} — {title_he}', bold=True, size=14, color=WHITE)
    rtl_run(p, f'  |  {title_en}', bold=False, size=10, color=(180,200,255))
    return p

def add_sub_heading(doc, number, title):
    """Sub-category heading"""
    p = doc.add_paragraph()
    make_rtl(p)
    shade_para(p, 'E8EEF8')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    set_para_border_bottom(p, '2C5594', '6')
    rtl_run(p, f'{number}  {title}', bold=True, size=12, color=NAVY)
    return p

def add_feature(doc, num, text):
    """Feature line with RTL formatting"""
    p = doc.add_paragraph()
    make_rtl(p)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.right_indent = Cm(0.3)
    rtl_run(p, f'{num}  ', bold=True, size=10, color=GOLD)
    rtl_run(p, text, bold=False, size=10, color=DARK)
    return p

def add_checkbox_line(doc, indent=True):
    """The [ ] כן  [ ] לא  [ ] הערות: ______ line"""
    p = doc.add_paragraph()
    make_rtl(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    if indent:
        p.paragraph_format.right_indent = Cm(0.6)
    rtl_run(p, '[ ] כן     [ ] לא     [ ] הערות:  ', bold=False, size=9.5, color=GREY)
    rtl_run(p, '________________________________', bold=False, size=9.5, color=(180,180,180))
    set_para_border_bottom(p, 'E0E0E0', '2')
    return p

def add_open_question(doc, num, text):
    p = doc.add_paragraph()
    make_rtl(p)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    rtl_run(p, f'{num}  ', bold=True, size=10, color=GOLD)
    rtl_run(p, text, bold=False, size=10, color=DARK)
    p2 = doc.add_paragraph()
    make_rtl(p2)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(10)
    rtl_run(p2, 'תשובה:  ', bold=False, size=9.5, color=GREY)
    rtl_run(p2, '___________________________________________________________', bold=False, size=9.5, color=(180,180,180))
    set_para_border_bottom(p2, 'B8860B', '3')
    return p

def add_body(doc, text, bold=False, size=11, color=DARK, italic=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        make_rtl(p)
    rtl_run(p, text, bold=bold, size=size, color=color, italic=italic)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    set_para_border_bottom(p, 'B8860B', '6')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

# =========================================================
# COVER PAGE
# =========================================================
doc.add_paragraph()
doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run('מסמך גילוי פיצ׳רים')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(*NAVY)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run('מיפוי צרכים ואפיון מערכת')
run2.bold = True
run2.font.size = Pt(20)
run2.font.color.rgb = RGBColor(*GOLD)

doc.add_paragraph()

p_en = doc.add_paragraph()
p_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p_en.add_run('Feature Discovery & Needs Assessment Checklist')
run3.bold = False
run3.font.size = Pt(14)
run3.font.color.rgb = RGBColor(*GREY)
run3.font.italic = True

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

cover_data = [
    ('מוגש ל:', '[שם הסאלון]'),
    ('מוגש על ידי:', '[שם צוות הפיתוח]'),
    ('תאריך:', '15 באפריל 2026'),
    ('גרסה:', '1.0 — סודי ומיועד לשימוש פנימי'),
]
for label, val in cover_data:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}  ')
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(*GREY)
    r2 = p.add_run(val)
    r2.bold = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(*NAVY)

doc.add_page_break()

# =========================================================
# WELCOME LETTER
# =========================================================
add_section_heading(doc, 'פתיחה', 'מכתב ברוכים הבאים', 'Welcome Letter')
doc.add_paragraph()

add_body(doc, 'לכבוד בעל העסק,', bold=True, size=12, color=NAVY)
doc.add_paragraph()
add_body(doc, (
    'אנו שמחים ומתרגשים להיפגש אתכם.'
), size=11)
add_body(doc, (
    'המסמך שלפניכם אינו רשימת מחירים, ואינו הצעת מכירה — זהו מפגש חלומות. '
    'מטרתו היחידה היא לאפשר לנו להבין לעומק את הוויזיה שלכם, את הצרכים העסקיים, '
    'ואת הפוטנציאל הבלתי מוגבל של הפלטפורמה הטכנולוגית שנבנה יחד.'
), size=11)
add_body(doc, (
    'אנו צוות פיתוח בוטיק — שלושה מפתחים בכירים עם ניסיון מצטבר של מעל 20 שנה — שמאמינים '
    'שסאלון ספרות מודרני ומוצלח כמו שלכם ראוי לפתרון טכנולוגי שלא ייפול מ-Nike, Spotify, '
    'או כל מותג גלובלי אחר. כל פיצ׳ר ברשימה זו ניתן לבנייה. אין כאן גבולות.'
), size=11)
add_body(doc, (
    'עברו על הרשימה בקצב שלכם, סמנו את מה שמדבר אליכם, הוסיפו הערות חופשיות — '
    'ומהשיחה הזו ניבנה יחד את המערכת המושלמת עבורכם.'
), size=11, italic=True, color=GREY)
doc.add_paragraph()
add_divider(doc)

doc.add_page_break()

# =========================================================
# SECTION A — Frontend & Brand
# =========================================================
add_section_heading(doc, 'א׳', 'נוכחות דיגיטלית, אתר ומיתוג', 'Frontend, Website & Brand')
doc.add_paragraph()

add_sub_heading(doc, '1.', 'עמוד הבית (Landing Page)')
features_1 = [
    ('1.1', 'עיצוב Landing Page מרהיב ומותאם אישית לזהות המותג של הסאלון'),
    ('1.2', 'אנימציות כניסה ואפקטים ויזואליים פרמיום (Scroll animations, Parallax)'),
    ('1.3', 'וידאו רקע בעמוד הבית המציג את האווירה הייחודית של הסאלון'),
    ('1.4', 'פקד קביעת תור מוטמע ישירות בעמוד הבית — ללא מעבר לדף נפרד'),
    ('1.5', 'כפתור "קבע תור עכשיו" קבוע (Sticky CTA) שנגרר עם גלילת העמוד'),
    ('1.6', 'אינטגרציה מלאה עם Google Maps + כפתור ניווט ישיר לסאלון'),
    ('1.7', 'פיד חי של ביקורות Google עם דירוג עדכני'),
]
for num, text in features_1:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '2.', 'גלריה אינטראקטיבית ותיק עבודות')
features_2 = [
    ('2.1', 'גלריית לפני/אחרי (Before & After) אינטראקטיבית עם Slider השוואה'),
    ('2.2', 'גלריה מסוננת לפי סוג שירות (תספורות, זקן, צביעה, תסרוקות אירוע)'),
    ('2.3', 'אינטגרציה אוטומטית עם Instagram — תמונות חדשות מופיעות בגלריה ישירות'),
    ('2.4', 'גלריה וידאו — Reels ו-TikToks מוטמעים בדף בצורה מעוצבת'),
    ('2.5', 'שיתוף תמונה מהגלריה ישירות לוואטסאפ/אינסטגרם'),
]
for num, text in features_2:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '3.', 'פרופילי הצוות')
features_3 = [
    ('3.1', 'דף ייעודי לכל ספר עם תמונה, ביו מקצועי, וסגנונות התמחות'),
    ('3.2', 'לחצן "קבע תור עם [שם הספר]" ישירות מדף הפרופיל'),
    ('3.3', 'תיק עבודות אישי לכל ספר עם גלריה נפרדת'),
    ('3.4', 'הצגת דירוג ממוצע וביקורות לקוחות לכל ספר בנפרד'),
    ('3.5', '"הספר הזמין עכשיו" — תג דינמי שמראה אם הספר פנוי ברגע זה'),
]
for num, text in features_3:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '4.', 'קורסי ספרות — Funnel הרשמה')
features_4 = [
    ('4.1', 'דף נחיתה ייעודי לקורסי הספרות המקצועיים (Sales Page מלא)'),
    ('4.2', 'טופס הרשמה לקורס עם תשלום מקדמה מאובטח'),
    ('4.3', 'ניהול מחזורי קורס — תאריכים, מקומות פנויים, רשימת המתנה'),
    ('4.4', 'שליחה אוטומטית של חומרי לימוד ומצגות לנרשמים'),
    ('4.5', 'תעודת קורס דיגיטלית (Certificate) שנשלחת אוטומטית עם סיום'),
    ('4.6', 'פורטל תלמיד ייעודי לגישה לחומרי לימוד, סרטוני הדרכה, ולוח שיעורים'),
    ('4.7', 'Upsell אוטומטי — הצעה לרכוש "ציוד מתחילים" מהחנות בעת ההרשמה'),
]
for num, text in features_4:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '5.', 'SEO, מהירות ונגישות')
features_5 = [
    ('5.1', 'אופטימיזציה מלאה למנועי חיפוש (Technical SEO) לדירוג "ספר בעיר X"'),
    ('5.2', 'Schema Markup לעסקים מקומיים (Local Business Schema)'),
    ('5.3', 'ציון Google PageSpeed של 90+ — טעינה מהירה מאוד'),
    ('5.4', 'עמידה בתקני נגישות WCAG 2.1 (חוק נגישות האתרים בישראל)'),
    ('5.5', 'ממשק מלא בעברית (RTL) עם אפשרות לאנגלית/ערבית'),
]
for num, text in features_5:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

doc.add_page_break()

# =========================================================
# SECTION B — Booking Engine
# =========================================================
add_section_heading(doc, 'ב׳', 'מנוע תורים מתקדם', 'Advanced Booking Engine')
doc.add_paragraph()

add_sub_heading(doc, '6.', 'קביעת תורים — יסודות')
features_6 = [
    ('6.1', 'קביעת תורים עצמאית 24/7 ללא צורך בשיחת טלפון'),
    ('6.2', 'בחירת ספר ספציפי או "כל ספר זמין"'),
    ('6.3', 'בחירת שירות מרשימה מפורטת עם מחירים ומשכי זמן'),
    ('6.4', 'הצגת חלונות זמן פנויים בזמן אמת (Real-time availability)'),
    ('6.5', 'תזכורות אוטומטיות ב-SMS/WhatsApp 24 שעות ו-2 שעות לפני התור'),
    ('6.6', 'אישור תור אוטומטי + ביטול/שינוי עצמאי על ידי הלקוח'),
    ('6.7', 'סנכרון אוטומטי עם Google Calendar ו-Apple Calendar של הספר'),
]
for num, text in features_6:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '7.', 'לוגיקה חכמה ואוטומציה')
features_7 = [
    ('7.1', 'אלגוריתם "מילוי פערים" — מציע חלונות זמן שמונעים חורים בלוח'),
    ('7.2', 'מניעת חפיפות אוטומטית בין ספרים ומשאבים (כסאות, מקלחות)'),
    ('7.3', '"זמן חיץ" אוטומטי בין תורים לניקוי/הכנה (Buffer Time)'),
    ('7.4', 'תורים מרובי-שירות בהזמנה אחת (תספורת + זקן + עיסוי קרקפת)'),
    ('7.5', 'זיהוי לקוח חוזר — הצעת ספר/שירות מועדף אוטומטית'),
    ('7.6', '"תור מהיר" — לקוח VIP יכול לתפוס תור פנוי אחרון בלחיצה אחת'),
]
for num, text in features_7:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '8.', 'רשימת המתנה ומקדמות')
features_8 = [
    ('8.1', 'רשימת המתנה חכמה — ביטול תור מייצר התראה לכל הממתינים'),
    ('8.2', '"תפוס את הביטול" — הודעת Push/WhatsApp על תור שהתפנה'),
    ('8.3', 'דרישת מקדמה (Deposit) בקביעת תור לצמצום No-Shows'),
    ('8.4', 'מדיניות ביטול גמישה — ביטול מעל 24 שעות = החזר; ביטול מאוחר = אין החזר'),
    ('8.5', 'הצעת תור קבוע שבועי/דו-שבועי ללקוחות חוזרים'),
    ('8.6', 'תזמון תור עתידי חכם — המערכת מציעה תאריך חזרה לפי תדירות ביקור ממוצעת'),
]
for num, text in features_8:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '9.', 'Walk-In וניהול תורים בזמן אמת')
features_9 = [
    ('9.1', 'מסך "תור חי" (Live Queue) לצג בסאלון — לקוחות רואים מה מספרם'),
    ('9.2', 'Check-In דיגיטלי — לקוח מגיע לסאלון וסורק QR לאישור הגעה'),
    ('9.3', 'ניהול Walk-Ins — הוספת לקוח אורח ללוח התורים בלחיצה אחת'),
    ('9.4', 'זמן המתנה משוער — הודעה ללקוח "התורך יתחיל בעוד ~20 דקות"'),
]
for num, text in features_9:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

doc.add_page_break()

# =========================================================
# SECTION C — POS & E-Commerce
# =========================================================
add_section_heading(doc, 'ג׳', 'קופה, חנות אונליין וניהול מלאי', 'POS, E-Commerce & Inventory')
doc.add_paragraph()

add_sub_heading(doc, '10.', 'חנות מוצרים אונליין')
features_10 = [
    ('10.1', 'חנות אונליין מלאה למכירת מוצרי שיער וזקן (שמפו, שמנים, מסרקים)'),
    ('10.2', 'עמודי מוצר עם תמונות מרובות, תיאור, ביקורות ודירוגי לקוחות'),
    ('10.3', 'מנוע המלצות חכם — "לקוחות שקנו X קנו גם Y"'),
    ('10.4', 'Bundle Deals — "קנה שמפו + בלסם ב-15% הנחה"'),
    ('10.5', 'Subscription מוצרים — שליחה חודשית קבועה של מוצר מועדף'),
    ('10.6', 'קוד קופון ומבצעים — אחוז הנחה / סכום קבוע / משלוח חינם'),
    ('10.7', '"הוסף מוצר לתור" — לקוח קונה אונליין ומקבל בסאלון'),
    ('10.8', 'Gift Card דיגיטלי — לרכישה ושליחה ישירה לוואטסאפ'),
]
for num, text in features_10:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '11.', 'ניהול מלאי')
features_11 = [
    ('11.1', 'מעקב מלאי שוטף — כמויות מתעדכנות אוטומטית עם כל מכירה/שימוש'),
    ('11.2', 'הבחנה בין מלאי "לשימוש בסאלון" לבין מלאי "למכירה ללקוח"'),
    ('11.3', 'התראת "מלאי נמוך" — Push + מייל כשמוצר מתקרב לאזילה'),
    ('11.4', 'הפקת "רשימת קניות" אוטומטית לספקים בלחיצה אחת'),
    ('11.5', 'ניהול עלות מול מחיר מכירה — רווחיות לכל מוצר'),
    ('11.6', 'ברקוד / QR Code לכל מוצר לסריקה מהירה בקופה הפיזית'),
    ('11.7', 'מעקב תפוגה (Expiry Date) למוצרים בעלי תאריך תפוגה'),
]
for num, text in features_11:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '12.', 'קופה פיזית (POS)')
features_12 = [
    ('12.1', 'ממשק קופה דיגיטלי (Tablet POS) לשימוש בסאלון עצמו'),
    ('12.2', 'תמיכה בתשלום מזומן, אשראי, Apple Pay, Google Pay'),
    ('12.3', 'אינטגרציה עם מסופי סליקה ישראליים (Sumit Pay, Cardcom, Tranzila)'),
    ('12.4', 'פיצול תשלום בין כמה אנשים (Split Bill) — שימושי לקבוצות'),
    ('12.5', 'קבלה דיגיטלית — שליחה אוטומטית במייל/וואטסאפ לאחר כל עסקה'),
    ('12.6', 'חשבוניות ירוקות (דיגיטליות) תואמות רשות המסים'),
    ('12.7', 'ניהול קופת מזומן — פתיחת/סגירת יום, ספירת קופה, דיווחים'),
    ('12.8', 'הנחת צוות — כפתור "עובד" שמחיל הנחה אוטומטית לבני הצוות'),
]
for num, text in features_12:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

doc.add_page_break()

# =========================================================
# SECTION D — CRM & Marketing
# =========================================================
add_section_heading(doc, 'ד׳', 'CRM ואוטומציה שיווקית', 'CRM & Marketing Automation')
doc.add_paragraph()

add_sub_heading(doc, '13.', 'ניהול קשרי לקוחות (CRM)')
features_13 = [
    ('13.1', 'תיק לקוח מלא: שם, טלפון, מייל, תאריך לידה, העדפות, הערות'),
    ('13.2', 'היסטוריית ביקורים מלאה — תאריכים, ספר, שירות, מחיר'),
    ('13.3', 'היסטוריית רכישות — כל מוצר שנקנה, מתי ובאיזה מחיר'),
    ('13.4', 'תיוג לקוחות (Tags) — VIP, לקוח בעייתי, סטודנט, ביקור ראשון'),
    ('13.5', 'ציון "ערך לקוח" (LTV — Lifetime Value) אוטומטי לכל לקוח'),
    ('13.6', 'הערות פרטיות של הספר על הלקוח (גלויות לצוות בלבד)'),
]
for num, text in features_13:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '14.', 'תקשורת ומסרים אוטומטיים')
features_14 = [
    ('14.1', 'שליחת SMS אוטומטי — אישור תור, תזכורת, ביטול'),
    ('14.2', 'WhatsApp Business API רשמי — אישורים, תזכורות, מסרים שיווקיים'),
    ('14.3', 'שליחת מסרים מקובצת (Bulk Message) למגזרי לקוחות ספציפיים'),
    ('14.4', 'תבניות הודעה מוכנות + עורך תבניות ויזואלי לבעל העסק'),
    ('14.5', 'שליחת מייל שיווקי (Newsletter) עם עורך Drag & Drop מובנה'),
    ('14.6', 'ניהול Opt-in / Opt-out לפי חוק הספאם הישראלי'),
]
for num, text in features_14:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '15.', 'קמפיינים אוטומטיים')
features_15 = [
    ('15.1', 'קמפיין יום הולדת — ברכה אישית + קופון הנחה ביום ההולדת'),
    ('15.2', 'קמפיין "אנחנו מתגעגעים" — לקוח שלא ביקר מעל X ימים'),
    ('15.3', '"זמן חזרה אופטימלי" — המערכת מחשבת מתי הלקוח "אמור" לחזור ומזכירה לו'),
    ('15.4', 'קמפיין "ביקור ראשון" — סדרת הודעות ברוכים הבאים (Welcome Series)'),
    ('15.5', 'קמפיין "לקוח VIP" — הזמנה לאירועים/השקות/מבצעים סגורים'),
    ('15.6', 'קמפיין "שגריר" — לקוח ממליץ לחבר ומקבל הנחה'),
    ('15.7', 'קמפיין עונתי — ראש השנה, חגים, קיץ/חורף, מבצעים אוטומטיים'),
]
for num, text in features_15:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '16.', 'מנויים ותוכניות VIP')
features_16 = [
    ('16.1', 'מנוי חודשי VIP — X תספורות בחודש במחיר מוזל (Membership)'),
    ('16.2', 'חבילות שירות מראש — רכישת 5 תספורות מראש בהנחה'),
    ('16.3', 'ניהול מנויים — חידוש אוטומטי, ביטול, השהיה'),
    ('16.4', 'הטבות VIP: אירועים, מוצרי ברכה, עדיפות בתורים'),
]
for num, text in features_16:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

doc.add_page_break()

# =========================================================
# SECTION E — Customer Portal
# =========================================================
add_section_heading(doc, 'ה׳', 'פורטל לקוח ואפליקציית Web', 'Customer Portal & Loyalty App')
doc.add_paragraph()

add_sub_heading(doc, '17.', 'פרופיל לקוח אישי')
features_17 = [
    ('17.1', 'אזור אישי מאובטח לכל לקוח (Login עם מייל/טלפון/Google)'),
    ('17.2', 'תצוגת היסטוריית ביקורים, שירותים, ומחירים'),
    ('17.3', 'שינוי/ביטול תור קיים ישירות מהפורטל'),
    ('17.4', 'ניהול פרטים אישיים, מספר טלפון, מייל, וכתובת'),
    ('17.5', 'שמירת כרטיס אשראי לתשלום מהיר בביקור הבא (Tokenization מאובטח)'),
]
for num, text in features_17:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '18.', '"תיק תסרוקת" דיגיטלי')
features_18 = [
    ('18.1', 'שמירת תמונות לפני/אחרי לכל ביקור, גלויות ללקוח בפורטל'),
    ('18.2', 'תיעוד "הנוסחה שלי" — אחוז צבע, טון, טכניקה — גלוי ללקוח ולספר'),
    ('18.3', '"הזמן את אותו הדבר שוב" — קביעת תור עם שירותים זהים לביקור קודם'),
    ('18.4', 'הערות אלרגיה ומוצרים לא מתאימים — מסומנים בולט לספר'),
]
for num, text in features_18:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '19.', 'נאמנות ומסלול נקודות')
features_19 = [
    ('19.1', 'מסלול נקודות — כל ₪ שהוצא = X נקודות; צבירה למימוש'),
    ('19.2', '"כרטיסיית ניקוב" דיגיטלית — 10 תספורות, האחת חינם'),
    ('19.3', 'לוח תוצאות (Leaderboard) — לקוחות הנאמנים ביותר החודש'),
    ('19.4', 'Referral Program — קישור ייחודי; המלצה = נקודות לשניהם'),
    ('19.5', 'Badges/הישגים — "לקוח שנה", "10 ביקורים", "רכישה ראשונה"'),
]
for num, text in features_19:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '20.', 'אפליקציה (PWA / Native App)')
features_20 = [
    ('20.1', 'Progressive Web App (PWA) — האתר "מותקן" בסמארטפון כאפליקציה'),
    ('20.2', 'התראות Push דרך האפליקציה (אישור תור, מבצע, ביטול)'),
    ('20.3', 'אפליקציה Native מלאה (iOS + Android) בחנות האפליקציות'),
]
for num, text in features_20:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

doc.add_page_break()

# =========================================================
# SECTION F — In-Shop / Barber-Tech
# =========================================================
add_section_heading(doc, 'ו׳', 'חוויית הסאלון — "ברבר-טק"', 'In-Shop Barber-Tech Experience')
doc.add_paragraph()

add_sub_heading(doc, '21.', 'בר השתייה — הזמנה דיגיטלית')
features_21 = [
    ('21.1', 'תפריט בר דיגיטלי — לקוח סורק QR מהכיסא ורואה את כל המשקאות'),
    ('21.2', 'הזמנת משקה דרך הסמארטפון בזמן שהספר עובד — ללא קום מהכיסא'),
    ('21.3', 'הוספת ההזמנה לחשבון התור ותשלום אחד מאוחד בסוף'),
    ('21.4', 'ניהול מלאי הבר — בקבוקים, קפסולות קפה, מתכלים'),
    ('21.5', '"Happy Hour בבר" — הנחה אוטומטית על משקאות בשעות שקטות'),
]
for num, text in features_21:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '22.', 'חוויה מוזיקלית חברתית')
features_22 = [
    ('22.1', 'הצבעה על מוסיקה — לקוח מצביע על הרצועה הבאה בפלייליסט Spotify'),
    ('22.2', '"הוסף שיר" — לקוח מציע שיר לנגן (בכפוף לאישור)'),
    ('22.3', 'פלייליסטים מותאמים לשעת היום (בוקר רגוע, ערב אנרגטי)'),
    ('22.4', 'שיתוף הפלייליסט של הסאלון לאחר הביקור — "המוזיקה מהסאלון אצלך בבית"'),
]
for num, text in features_22:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '23.', 'AR — מציאות רבודה לתסרוקות')
features_23 = [
    ('23.1', 'AR Hairstyle Try-On — לקוח מנסה תסרוקות שונות דרך המצלמה לפני ההחלטה'),
    ('23.2', 'AR זקן — ניסוי סגנונות זקן שונים על הפנים בזמן אמת'),
    ('23.3', 'שמירת ה"לוק" הנבחר כהמלצה לספר לפני תחילת הטיפול'),
    ('23.4', 'שיתוף תוצאת ה-AR ברשתות חברתיות — "כך ייראה הלוק החדש שלי"'),
]
for num, text in features_23:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '24.', 'מסכי סאלון ותצוגה דיגיטלית')
features_24 = [
    ('24.1', 'Digital Signage — מסכי טלוויזיה בסאלון עם תוכן מנוהל מהמערכת'),
    ('24.2', 'הצגת תור חי על מסך — לקוחות ממתינים רואים מצב התור'),
    ('24.3', 'הצגת מבצעים ומוצרים על מסכי הסאלון'),
    ('24.4', 'הצגת תיק עבודות / הישגי הצוות על מסכי הסאלון'),
]
for num, text in features_24:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '25.', 'Gamification וחוויה')
features_25 = [
    ('25.1', '"אתגר חודשי" — לקוח שמגיע X פעמים בחודש מקבל פרס'),
    ('25.2', 'Check-In Social — שיתוף ביקור לאינסטגרם/Facebook (בהרשאה)'),
    ('25.3', '"לקוח החודש" — תצוגה בסאלון ובאתר (בהסכמה)'),
    ('25.4', '"ספר מומלץ" — לקוח ממליץ על הספר שלו לחברים דרך קישור ייחודי'),
]
for num, text in features_25:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '26.', 'Wi-Fi ו-Smart Waiting')
features_26 = [
    ('26.1', 'Wi-Fi אורחים עם Captive Portal בעיצוב המותג + לכידת מייל/טלפון'),
    ('26.2', 'חיבור ל-Wi-Fi מעניק גישה אוטומטית לפורטל הלקוח'),
    ('26.3', 'סקר שביעות רצון קצר (שאלה אחת) שנשלח 30 דקות לאחר הביקור'),
]
for num, text in features_26:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

doc.add_page_break()

# =========================================================
# SECTION G — Staff & Management
# =========================================================
add_section_heading(doc, 'ז׳', 'לוח בקרה לצוות וניהול עסקי', 'Staff & Management Dashboard')
doc.add_paragraph()

add_sub_heading(doc, '27.', 'ניהול משמרות וצוות')
features_27 = [
    ('27.1', 'לוח שנה חכם לניהול משמרות — גרירה ושחרור, החלפת משמרות'),
    ('27.2', 'בקשת חופשה/מחלה דיגיטלית — עובד מגיש, מנהל מאשר'),
    ('27.3', 'שעון נוכחות דיגיטלי — Check-In/Out עם QR או PIN'),
    ('27.4', 'ניהול חוזים ומסמכי עובדים (שמירה בענן)'),
    ('27.5', 'הודעות פנימיות בין בעל העסק לצוות (Internal Messaging)'),
]
for num, text in features_27:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '28.', 'תגמולים ועמלות')
features_28 = [
    ('28.1', 'חישוב עמלה אוטומטי לכל ספר לפי הכנסות חודשיות'),
    ('28.2', 'מבנה עמלה גמיש — אחוז על שירותים, מוצרים, בונוס מעל יעד'),
    ('28.3', '"לוח תוצאות ספרים" — דירוג פנימי הכנסות/ביקורות/No-Shows (להנהלה בלבד)'),
    ('28.4', 'דוח שכר חודשי אוטומטי לכל עובד (שכר בסיס + עמלות)'),
    ('28.5', 'יעדים חודשיים לכל ספר עם מד התקדמות חזותי'),
]
for num, text in features_28:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '29.', 'אנליטיקה ובינה עסקית (BI)')
features_29 = [
    ('29.1', 'דאשבורד הכנסות: יומי, שבועי, חודשי, שנתי — עם גרפים'),
    ('29.2', 'ניתוח "שעות שיא" — באילו ימים/שעות הסאלון הכי עמוס'),
    ('29.3', 'שיעור No-Show לפי ספר, שירות, ויום בשבוע'),
    ('29.4', 'שיעור שימור לקוחות (Retention Rate) — כמה לקוחות חוזרים'),
    ('29.5', 'ניתוח רווחיות מוצרים — אילו פריטים מייצרים הכי הרבה רווח'),
    ('29.6', 'תחזית הכנסות (Revenue Forecast) — AI שמנבא הכנסה חודש קדימה'),
    ('29.7', 'דוחות ייצוא — Excel, PDF — לשיתוף עם רואה החשבון'),
    ('29.8', 'השוואת תקופות — "ינואר 2025 מול ינואר 2026"'),
]
for num, text in features_29:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '30.', 'ניהול מוניטין ובקרת איכות')
features_30 = [
    ('30.1', 'בקשה אוטומטית לדירוג Google לאחר כל ביקור'),
    ('30.2', 'ריכוז כל הביקורות (Google, Facebook) בדאשבורד אחד'),
    ('30.3', 'דירוג פנימי — לקוח מדרג ביקור ישירות במערכת'),
    ('30.4', 'התראה למנהל על ביקורת שלילית — לטיפול מהיר'),
]
for num, text in features_30:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

doc.add_page_break()

# =========================================================
# SECTION H — Security & Infrastructure
# =========================================================
add_section_heading(doc, 'ח׳', 'תשתית, אבטחה ותאימות', 'Security, Infrastructure & Compliance')
doc.add_paragraph()

add_sub_heading(doc, '31.', 'תשתית ענן')
features_31 = [
    ('31.1', 'פריסה על VPS בענן עם SLA מובטח של 99.9% Uptime'),
    ('31.2', 'סביבות נפרדות — Development, Staging, Production'),
    ('31.3', 'CI/CD Pipeline — עדכוני תוכנה אוטומטיים ללא Downtime'),
    ('31.4', 'יכולת Scale Up — הגדלת שרת בלחיצה אחת בזמן עומס'),
]
for num, text in features_31:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '32.', 'אבטחת מידע')
features_32 = [
    ('32.1', 'Cloudflare Pro — הגנת DDoS, CDN, WAF (חומת אש לאפליקציה)'),
    ('32.2', 'SSL/TLS מלא על כל הדומיינים ותת-הדומיינים'),
    ('32.3', 'הצפנת נתונים AES-256 — במנוחה ובמעבר'),
    ('32.4', 'אימות דו-שלבי (2FA) לכל כניסת מנהל/עובד'),
    ('32.5', 'ניטור אבטחה 24/7 + התראות על פעילות חשודה'),
    ('32.6', 'Penetration Testing — בדיקת חדירות תקופתית'),
]
for num, text in features_32:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '33.', 'גיבויים ושחזור')
features_33 = [
    ('33.1', 'גיבוי יומי אוטומטי של מסד הנתונים לאחסון גיאוגרפי נפרד'),
    ('33.2', 'שמירת 30 גיבויים אחורה — שחזור לכל נקודת זמן'),
    ('33.3', 'בדיקת שחזור רבעונית — וידוא שהגיבוי אכן עובד'),
    ('33.4', 'Recovery Time Objective (RTO) — שחזור תוך פחות מ-4 שעות'),
]
for num, text in features_33:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '34.', 'הרשאות ובקרת גישה (RBAC)')
features_34 = [
    ('34.1', 'תפקיד "בעל עסק" — גישה מלאה לכל המערכת'),
    ('34.2', 'תפקיד "ספר בכיר" — CRM, תורים, תיקי לקוחות; אין גישה לנתונים כספיים'),
    ('34.3', 'תפקיד "ספר זוטר" — גישה לתורים שלו בלבד'),
    ('34.4', 'תפקיד "קופאי/מנהל חנות" — קופה ומלאי; אין גישה לנתוני לקוחות'),
    ('34.5', 'תפקיד "מדריך קורסים" — גישה לפורטל תלמידים ומחזורי קורס'),
    ('34.6', 'יצירת תפקידים מותאמים אישית לפי הצורך'),
    ('34.7', 'לוג גישה — רישום מי היה במערכת, מה שינה, ומתי'),
]
for num, text in features_34:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '35.', 'תאימות חוקית')
features_35 = [
    ('35.1', 'תאימות GDPR ו/או חוק הגנת הפרטיות הישראלי'),
    ('35.2', 'מדיניות פרטיות ותנאי שימוש מותאמים לפלטפורמה'),
    ('35.3', '"שכח אותי" — מחיקת נתוני לקוח מלאה על פי דרישה (Right to Erasure)'),
    ('35.4', 'Cookie Banner תואם לרגולציה האירופית'),
]
for num, text in features_35:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

doc.add_page_break()

# =========================================================
# SECTION I — Integrations & Future
# =========================================================
add_section_heading(doc, 'ט׳', 'אינטגרציות חיצוניות ועתיד', 'Third-Party Integrations & AI Future')
doc.add_paragraph()

add_sub_heading(doc, '36.', 'אינטגרציות פלטפורמה')
features_36 = [
    ('36.1', 'אינטגרציה עם Google Analytics 4 + Google Tag Manager'),
    ('36.2', 'Meta Pixel (Facebook/Instagram) לקמפיינים ממומנים'),
    ('36.3', 'אינטגרציה עם TikTok Pixel'),
    ('36.4', 'אינטגרציה עם מערכת הנהלת חשבונות — iCount, חשבשבת, Priority'),
    ('36.5', 'אינטגרציה עם Wolt / Ten-Bis לניהול הזמנות בר'),
    ('36.6', 'אינטגרציה עם Zapier / Make — חיבור לאלפי כלים חיצוניים'),
]
for num, text in features_36:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

add_sub_heading(doc, '37.', 'AI ועתיד')
features_37 = [
    ('37.1', 'Chatbot AI מבוסס GPT-4 באתר — ענה על שאלות, קבע תורים, המלץ שירותים'),
    ('37.2', 'המלצות שירות חכמות — AI מנתח היסטוריה ומציע "הגיע זמן לצבע שיער"'),
    ('37.3', 'Dynamic Pricing — AI מגדיר מחיר גבוה בשיא ומחיר מוזל בשעות שקטות'),
    ('37.4', 'זיהוי פנים — מצלמה מזהה לקוח בכניסה ופותחת תיק לקוח אוטומטית'),
    ('37.5', 'Voice Booking — קביעת תור דרך שיחת קול לבוט חכם בוואטסאפ'),
    ('37.6', 'ניתוח תמונה AI — לקוח מעלה תמונת לוק רצוי, AI מתאים שירות מתאים'),
    ('37.7', 'מערכת BI מתקדמת עם תחזיות מכירה מבוססות Machine Learning'),
]
for num, text in features_37:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

doc.add_page_break()

# =========================================================
# SECTION — Open Questions
# =========================================================
add_section_heading(doc, 'י׳', 'שאלות פתוחות ומחשבות חופשיות', 'Open Questions & Vision')
doc.add_paragraph()

add_body(doc, 'לפני שנסיים — כמה שאלות שחשוב לנו לשמוע עליהן:', bold=True, size=12, color=NAVY)
doc.add_paragraph()

open_qs = [
    ('38.1', 'האם יש פיצ׳ר שחלמתם עליו ולא ראיתם אצל אף אחד? תפרטו:'),
    ('38.2', 'מה הכי מתסכל אתכם כיום במערכת/כלים שבהם אתם משתמשים?'),
    ('38.3', 'מה הדבר אחד שאם היה קורה אוטומטי — היה חוסך לכם הכי הרבה זמן?'),
]
for num, text in open_qs:
    add_open_question(doc, num, text)

# Checkbox questions
checkbox_qs = [
    ('38.4', 'האם אתם מתכננים להרחיב לסניפים נוספים בעתיד?'),
    ('38.5', 'האם יש שיתופי פעולה עסקיים עתידיים שהמערכת צריכה לתמוך בהם?'),
]
for num, text in checkbox_qs:
    add_feature(doc, num, text)
    add_checkbox_line(doc)

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

# Closing
add_body(doc, 'תודה רבה על הזמן והפתיחות.', bold=True, size=13, color=NAVY)
add_body(doc, (
    'מהשיחה הזו ניבנה יחד את הפלטפורמה שתהפוך את הסאלון שלכם לחוויה שאנשים מספרים עליה.'
), size=11, italic=True, color=GREY)

doc.add_paragraph()
add_body(doc, f'מסמך זה הוכן על ידי: [שם הצוות]   |   © 2026   |   סודי ומיועד לשימוש פנימי בלבד', size=9, color=GREY, italic=True)

# Save
output_path = r'c:\Users\OWNER\MathU\רשימת_פיצ׳רים_סאלון_ספרות_2026.docx'
doc.save(output_path)
print(f'Checklist document saved to: {output_path}')
