# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page setup (A4) ---
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# --- RTL helper ---
def make_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def rtl_run(paragraph, text, bold=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rtl_elem = OxmlElement('w:rtl')
    rPr.append(rtl_elem)
    return run

def add_heading(doc, text, level=1, color=(0,0,0)):
    p = doc.add_paragraph()
    make_rtl(p)
    sizes = {1: 22, 2: 16, 3: 13}
    rtl_run(p, text, bold=True, size=sizes.get(level, 12), color=color)
    return p

def add_body(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    make_rtl(p)
    rtl_run(p, text, bold=bold, size=size)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    make_rtl(p)
    rtl_run(p, text, size=size)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=10, align='right', bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    make_rtl(p)
    rtl_run(p, text, bold=bold, size=size)
    if bg:
        shade_cell(cell, bg)

GOLD   = (184, 134, 11)
DARK   = (30, 30, 30)
MEDIUM = (70, 70, 70)
ACCENT = (44, 85, 148)

# =========================================================
# COVER PAGE
# =========================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('הצעת מחיר מקצועית')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(*GOLD)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('מערכת ניהול SaaS מלאה לסאלון יופי')
run2.bold = True
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(*ACCENT)

doc.add_paragraph()

cover_info = [
    ('מסמך מס׳:', '2026-001'),
    ('תאריך:', '15 באפריל 2026'),
    ('מוגשת ל:', '[שם הסאלון]'),
    ('מוגשת על ידי:', '[שם חברת הפיתוח]'),
    ('תוקף ההצעה:', '30 יום ממועד הגשה'),
]
for label, val in cover_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}  {val}')
    run.font.size = Pt(13)

doc.add_page_break()

# =========================================================
# SECTION 1 — EXECUTIVE SUMMARY
# =========================================================
add_heading(doc, 'סעיף 1 — סיכום מנהלים (Executive Summary)', 1, ACCENT)

add_body(doc, 'לכבוד הנהלת הסאלון,', bold=True)
add_body(doc, (
    'אנו מגישים בפניכם הצעה לפיתוח ומימוש פתרון תוכנה מנוהל (SaaS) מלא ומותאם אישית עבור עסק הספרות/יופי שלכם — '
    'פתרון שיהפוך את אופן ניהול הסאלון, יחסוך שעות עבודה מדי שבוע, יגדיל את שימור הלקוחות, וייצג את המותג שלכם ברמה הגבוהה ביותר.'
))

add_body(doc, 'מי אנחנו?', bold=True)
add_body(doc, (
    'אנו צוות פיתוח בוטיק של שלושה מפתחים בכירים, עם ניסיון מצטבר של מעל 20 שנה בבניית מערכות SaaS עסקיות מורכבות, '
    'פתרונות תשלומים, ואוטומציה שיווקית לעסקים בינוניים וגדולים בישראל ובעולם. '
    'אנו לא חנות לפיתוח זול ומהיר — אנו שותפים טכנולוגיים אסטרטגיים המחויבים להצלחה ארוכת הטווח של הלקוח.'
))

add_body(doc, 'האתגר שאנו פותרים:', bold=True)
add_body(doc, (
    'עסקי יופי רבים מסתמכים על שילוב לא יעיל של יומנים ידניים, WhatsApp לא מנוהל, גיליונות Excel, וכלים זולים שאינם מדברים זה עם זה. '
    'התוצאה: תורים שנשכחים, הכנסות שאינן עוקבות, לקוחות שנוטשים, וצוות שמבזבז זמן על ניהול במקום על מקצוע.'
))

add_body(doc, 'הפתרון שלנו:', bold=True)
add_body(doc, (
    'מערכת אחת, משולבת לחלוטין — מאתר האינטרנט ועד לחשבונית הדיגיטלית — שתנהל עבורכם את כל מחזור חיי הלקוח: '
    'מהרגע שהוא מגלה אתכם ברשת, דרך קביעת התור, הגעתו לסאלון, ביצוע התשלום, ועד ההודעה האוטומטית שמזמינה אותו לחזור.'
))

doc.add_page_break()

# =========================================================
# SECTION 2 — TECHNICAL SCOPE
# =========================================================
add_heading(doc, 'סעיף 2 — היקף טכני ותיאור הרכיבים', 1, ACCENT)

modules = [
    ('2.1 — נוכחות דיגיטלית ופרונטאנד (Web Presence & Frontend)', [
        'עיצוב UI/UX מותאם אישית: מחקר מתחרים, מפת מסע לקוח, Wireframes, Mockups, Prototype.',
        'אתר Landing Page עם המרה גבוהה: הצגת שירותים, גלריית עבודות, ביקורות לקוחות, ופקד קביעת תור מוטמע.',
        'רישום דומיין: רישום שם הדומיין (.co.il / .com) וניהולו לשנה הראשונה.',
        'אחסון מנוהל: הגדרת VPS בענן (DigitalOcean / AWS / Hetzner) עם סביבת הפקה מאובטחת.',
        'אופטימיזציה למנועי חיפוש (SEO): meta tags, schema markup, מהירות טעינה, mobile-first.',
        'ממשק מלא בעברית (RTL), נגישות תקן WCAG 2.1.',
    ]),
    ('2.2 — מערכת הניהול המרכזית (Core System — Dashboard & App)', [
        'לוח בקרה לבעל העסק והצוות (Admin Dashboard): ממשק ניהול עשיר ומרוכז, נגיש מכל מכשיר (Responsive Web App).',
        'פורטל לקוחות (Customer Portal): ממשק נקי המאפשר ללקוח לראות היסטוריית ביקורים, לקבוע תורים, ולנהל פרטיו.',
    ]),
    ('2.3 — מערכת תורים ולוח זמנים חכם (Smart Scheduling & Calendar)', [
        'קביעת תורים אוטומטית 24/7 ישירות מהאתר או מהפורטל.',
        'זיהוי וניהול קונפליקטים: מניעת חפיפות בין עובדים ומשאבים בזמן אמת.',
        'ניהול זמינות צוות: שעות, חופשות, הפסקות.',
        'סנכרון דו-כיווני עם Google Calendar ו-Apple Calendar (iCal).',
        'הגדרת סוגי שירות, משכי זמן, ומחירים לכל עובד/ת בנפרד.',
    ]),
    ('2.4 — תקשורת ומסרים אוטומטיים (Communications & Notifications)', [
        'שילוב SMS עם ספקים ישראליים (Inforu / Vonage / Twilio) — עלות ~₪0.08–0.15 להודעה.',
        'WhatsApp Business API רשמי (Meta): אישורים, תזכורות, מבצעים — ללא סכנת חסימה.',
        'תבניות הודעה מוכנות לשימוש מיידי.',
        'ניהול הסכמות Opt-in/Opt-out בהתאם לחוק הספאם הישראלי.',
    ]),
    ('2.5 — חיוב ותשלומים (Billing & Payments)', [
        'שילוב עם שירות סליקה ישראלי: Cardcom / Tranzila / Meshulam.',
        'תמיכה בכרטיסי אשראי, Apple Pay, ו-Google Pay.',
        'אפשרות לדרוש מקדמה בקביעת תור להפחתת no-shows.',
        'חשבוניות ירוקות: הפקה אוטומטית תואמת רשות המסים עם שליחה במייל/WhatsApp (iCount / Chequepoint).',
        'ניהול החזרים (Refunds) ישירות מלוח הבקרה.',
    ]),
    ('2.6 — מדדים ואנליטיקה עסקית (Data & Analytics)', [
        'הכנסה יומית/שבועית/חודשית: גרפים ומגמות לאורך זמן.',
        'שירותים מובילים: אילו טיפולים הכי פופולריים ורווחיים.',
        'ביצועי עובדים: מספר תורים, הכנסות, ודירוג ממוצע.',
        'שיעור שימור לקוחות (Retention Rate) וניתוח no-shows.',
        'ייצוא דוחות ל-Excel/CSV.',
    ]),
    ('2.7 — תשתית ואבטחה (Infrastructure & Security)', [
        'VPS מנוהל בענן עם SLA של 99.9% uptime.',
        'PostgreSQL עם גיבויים יומיים לאחסון גיאוגרפי נפרד.',
        'SSL/TLS מלא (Let\'s Encrypt / Cloudflare).',
        'הצפנת נתונים AES-256 במנוחה ובמעבר.',
        'תאימות GDPR / חוק הפרטיות הישראלי כולל "Right to Erasure".',
        'הגנת DDoS דרך Cloudflare וניטור 24/7.',
    ]),
    ('2.8 — CRM לקוחות מתקדם [פיצ׳ר פרמיום]', [
        'תיק לקוח מלא: שם, טלפון, מייל, תאריך לידה, העדפות, ורשומות ביקור.',
        'נוסחאות צבע ועיצוב שיער: תיעוד מדויק לכל גוון, מוצר ומינון.',
        'הערות אלרגיות: מוצגות בולט בעת פתיחת התור.',
        'גלריית לפני/אחרי לכל לקוח/ה — לשימוש שיווקי ברשתות (באישור הלקוח).',
    ]),
    ('2.9 — ניהול מלאי (Inventory Management) [פיצ׳ר פרמיום]', [
        'מעקב מלאי שוטף: כמויות מתעדכנות אוטומטית עם כל שימוש.',
        'הבחנה בין מוצרים פנימיים ומכירה ישירה ללקוח.',
        'התראות מלאי נמוך אוטומטיות.',
        'ייצוא רשימת קניות לספקים.',
    ]),
    ('2.10 — ניהול הרשאות מבוסס תפקיד (RBAC) [פיצ׳ר פרמיום]', [
        'בעל/ת העסק (Admin): גישה מלאה לכל המערכת.',
        'עובד/ת בכיר/ה (Senior Stylist): CRM, תורים, נוסחאות.',
        'עובד/ת זוטר/ה (Junior Staff): גישה מוגבלת לתורים בלבד.',
        'גמישות להוספת תפקידים חדשים בעתיד.',
    ]),
    ('2.11 — אוטומציה שיווקית (Marketing Automation) [פיצ׳ר פרמיום]', [
        'הודעות יום הולדת אוטומטיות + קוד קופון אישי.',
        'קמפיין "אנחנו מתגעגעים" — לקוח שלא ביקר מעל 60 יום.',
        'סגמנטציה: VIP, לקוחות חדשים, לקוחות עם ביטולים.',
        'מעקב ביצועי קמפיינים: אחוזי פתיחה, קליקים, המרות.',
    ]),
]

for title, bullets in modules:
    add_heading(doc, title, 2, ACCENT)
    for b in bullets:
        add_bullet(doc, b)
    doc.add_paragraph()

doc.add_page_break()

# =========================================================
# SECTION 3 — PRICING
# =========================================================
add_heading(doc, 'סעיף 3 — פירוט תמחור (Pricing Breakdown)', 1, ACCENT)
add_body(doc, (
    'המחירים משקפים שיעורי שוק עדכניים לשנת 2026 לפיתוח SaaS מותאם אישית בישראל (~₪250–350/שעה לפיתוח בכיר). '
    'הצעה זו אינה מחיר פרילנסר בסיסי — זהו תמחור של פתרון ארגוני-עסקי מלא עם SLA ברור ואחריות מלאה.'
))
doc.add_paragraph()

# --- Development table ---
add_heading(doc, 'מודל א׳ — עלות פיתוח והקמה (חד-פעמי)', 2, ACCENT)

dev_rows = [
    ('שלב 1 — עיצוב (Design)', 'מחקר מתחרים, מפת מסע לקוח, Wireframes, Mockups, Prototype', '₪12,500'),
    ('שלב 2א — פיתוח בסיס', 'תשתית Backend, מסד נתונים, Authentication, RBAC', '₪22,000'),
    ('שלב 2ב — פיתוח Frontend', 'אתר Landing Page, Admin Dashboard, Customer Portal', '₪25,000'),
    ('שלב 2ג — מערכת תורים', 'מנוע Scheduling, Google/Apple Calendar sync', '₪18,000'),
    ('שלב 2ד — תקשורת ומסרים', 'שילוב SMS, WhatsApp Business API', '₪10,000'),
    ('שלב 2ה — תשלומים וחשבוניות', 'שילוב סליקה, Apple/Google Pay, חשבוניות ירוקות', '₪12,500'),
    ('שלב 2ו — CRM + מלאי + אנליטיקה', 'פיצ׳רים פרמיום: CRM, Inventory, Analytics Dashboard', '₪22,000'),
    ('שלב 2ז — אוטומציה שיווקית', 'Birthday/Win-back campaigns, segmentation, tracking', '₪9,000'),
    ('שלב 3 — בדיקות QA', 'Functional, Security, Performance, UAT', '₪8,000'),
    ('שלב 4 — השקה ותשתית', 'הגדרת VPS, CI/CD, DNS, SSL, גיבויים, ניטור', '₪6,000'),
    ('שלב 5 — הדרכה ואונבורדינג', 'סשן הדרכה לבעל/ת העסק ולצוות + מסמך הדרכה', '₪3,500'),
]

table1 = doc.add_table(rows=1, cols=3)
table1.style = 'Table Grid'
hdr = table1.rows[0].cells
headers = ['עלות (₪)', 'תיאור', 'שלב / משימה']
for i, h in enumerate(headers):
    set_cell_text(hdr[i], h, bold=True, size=10, bg='2C5594')
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

for idx, (stage, desc, price) in enumerate(dev_rows):
    row = table1.add_row().cells
    bg = 'F2F6FF' if idx % 2 == 0 else 'FFFFFF'
    set_cell_text(row[0], price, bold=True, size=10, bg=bg)
    set_cell_text(row[1], desc, size=10, bg=bg)
    set_cell_text(row[2], stage, bold=False, size=10, bg=bg)

total_row = table1.add_row().cells
set_cell_text(total_row[0], '₪148,500', bold=True, size=11, bg='B8860B')
total_row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
set_cell_text(total_row[1], '', size=10, bg='B8860B')
set_cell_text(total_row[2], 'סה״כ השקעה חד-פעמית', bold=True, size=11, bg='B8860B')
total_row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

doc.add_paragraph()
add_body(doc, 'מבנה תשלום מומלץ למיילסטונים:', bold=True)
milestones = [
    '40% עם חתימת ההסכם — ₪59,400',
    '30% עם אישור עיצובים וסיום שלב 2 Backend — ₪44,550',
    '20% עם סיום שלב 2 Frontend + אינטגרציות — ₪29,700',
    '10% עם השקה ואישור סופי — ₪14,850',
]
for m in milestones:
    add_bullet(doc, m)

doc.add_paragraph()

# --- Monthly retainer table ---
add_heading(doc, 'מודל ב׳ — ריטיינר SaaS ותפעול שוטף (חודשי/שנתי)', 2, ACCENT)

add_body(doc, 'עלויות תשתית ושירות (בשליטתנו):', bold=True)
infra_rows = [
    ('VPS Cloud Hosting', 'שרת מנוהל (DigitalOcean/Hetzner) 4vCPU / 8GB RAM', '₪180'),
    ('גיבויים + Storage', 'גיבוי יומי אוטומטי לאחסון גיאוגרפי נפרד (S3/Spaces)', '₪45'),
    ('Cloudflare Pro', 'הגנת DDoS, CDN, שיפור ביצועים', '₪90'),
    ('ניטור ואלרטים', 'Uptime monitoring, Error tracking (Sentry/BetterUptime)', '₪40'),
    ('אחזקת דומיין', '₪90/שנה ÷ 12', '₪8'),
    ('תמיכה טכנית + SLA', 'תגובה עד 24 שעות לבאגים קריטיים, עדכוני אבטחה, תיקונים', '₪950'),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, h in enumerate(['עלות חודשית (₪)', 'תיאור', 'רכיב']):
    set_cell_text(hdr2[i], h, bold=True, size=10, bg='2C5594')
    hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

for idx, (comp, desc, price) in enumerate(infra_rows):
    row = table2.add_row().cells
    bg = 'F2F6FF' if idx % 2 == 0 else 'FFFFFF'
    set_cell_text(row[0], price, bold=True, size=10, bg=bg)
    set_cell_text(row[1], desc, size=10, bg=bg)
    set_cell_text(row[2], comp, size=10, bg=bg)

total2 = table2.add_row().cells
set_cell_text(total2[0], '₪1,313 / חודש', bold=True, size=11, bg='B8860B')
total2[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
set_cell_text(total2[1], '', bg='B8860B')
set_cell_text(total2[2], 'סה״כ תשתית ושירות', bold=True, size=11, bg='B8860B')
total2[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

doc.add_paragraph()
add_body(doc, 'עלויות APIs של צד שלישי (הערכה ל-150–300 תורים/חודש):', bold=True)

api_rows = [
    ('SMS (Inforu / Vonage)', '~300 הודעות/חודש × ₪0.12', '₪36–55'),
    ('WhatsApp Business API', '~400 הודעות/חודש × ₪0.15 (ממוצע template)', '₪55–90'),
    ('שירות סליקה', 'עמלת 0.8%–1.5% + דמי קבועים (מחזור ~₪20,000)', '₪80–200'),
    ('חשבוניות (iCount)', 'תוכנית בסיס', '₪49'),
]

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
for i, h in enumerate(['עלות חודשית (₪)', 'הערכת שימוש', 'שירות']):
    set_cell_text(hdr3[i], h, bold=True, size=10, bg='2C5594')
    hdr3[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

for idx, (svc, usage, price) in enumerate(api_rows):
    row = table3.add_row().cells
    bg = 'F2F6FF' if idx % 2 == 0 else 'FFFFFF'
    set_cell_text(row[0], price, bold=True, size=10, bg=bg)
    set_cell_text(row[1], usage, size=10, bg=bg)
    set_cell_text(row[2], svc, size=10, bg=bg)

total3 = table3.add_row().cells
set_cell_text(total3[0], '₪220–394 / חודש', bold=True, size=11, bg='B8860B')
total3[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
set_cell_text(total3[1], '', bg='B8860B')
set_cell_text(total3[2], 'סה״כ APIs (הערכה)', bold=True, size=11, bg='B8860B')
total3[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

doc.add_paragraph()

# --- Grand total table ---
add_heading(doc, 'ריכוז עלויות תפעול חודשיות', 2, ACCENT)

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
for i, h in enumerate(['עלות שנתית', 'עלות חודשית', '']):
    set_cell_text(hdr4[i], h, bold=True, size=10, bg='2C5594')
    hdr4[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

grand_rows = [
    ('תשתית ושירות', '₪1,313', '₪15,756'),
    ('APIs צד שלישי (ממוצע)', '₪307', '₪3,684'),
]
for idx, (label, mo, yr) in enumerate(grand_rows):
    row = table4.add_row().cells
    bg = 'F2F6FF' if idx % 2 == 0 else 'FFFFFF'
    set_cell_text(row[0], yr, bold=True, size=10, bg=bg)
    set_cell_text(row[1], mo, bold=True, size=10, bg=bg)
    set_cell_text(row[2], label, size=10, bg=bg)

total4 = table4.add_row().cells
set_cell_text(total4[0], '₪19,440 / שנה', bold=True, size=11, bg='B8860B')
total4[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
set_cell_text(total4[1], '₪1,620 / חודש', bold=True, size=11, bg='B8860B')
total4[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
set_cell_text(total4[2], 'סה״כ ריטיינר', bold=True, size=11, bg='B8860B')
total4[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

doc.add_paragraph()
add_body(doc, 'מסלול שנתי מומלץ: תשלום שנתי מראש יקנה הנחה של 10% — ₪17,496 / שנה (חיסכון של ~₪1,944).', bold=True)

doc.add_page_break()

# =========================================================
# SECTION 4 — TOTAL INVESTMENT
# =========================================================
add_heading(doc, 'סעיף 4 — סיכום השקעה כולל', 1, ACCENT)

table5 = doc.add_table(rows=1, cols=2)
table5.style = 'Table Grid'
hdr5 = table5.rows[0].cells
for i, h in enumerate(['עלות', 'שלב']):
    set_cell_text(hdr5[i], h, bold=True, size=11, bg='2C5594')
    hdr5[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

summary_rows = [
    ('פיתוח והקמה (חד-פעמי)', '₪148,500'),
    ('ריטיינר שנה ראשונה (מסלול שנתי)', '₪17,496'),
]
for idx, (label, price) in enumerate(summary_rows):
    row = table5.add_row().cells
    bg = 'F2F6FF' if idx % 2 == 0 else 'FFFFFF'
    set_cell_text(row[0], price, bold=True, size=11, bg=bg)
    set_cell_text(row[1], label, size=11, bg=bg)

total5 = table5.add_row().cells
set_cell_text(total5[0], '₪165,996', bold=True, size=13, bg='1E3A5F')
total5[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,215,0)
set_cell_text(total5[1], 'סה״כ השקעה לשנה הראשונה', bold=True, size=13, bg='1E3A5F')
total5[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,215,0)

doc.add_page_break()

# =========================================================
# SECTION 5 — TERMS & SLA
# =========================================================
add_heading(doc, 'סעיף 5 — תנאים והתחייבויות (Terms & SLA)', 1, ACCENT)

add_heading(doc, '5.1 — לו"ז פיתוח משוער', 2, ACCENT)
timeline_rows = [
    ('שלב 1 — עיצוב ואישור', 'שבועות 1–3'),
    ('שלב 2 — פיתוח (כולל)', 'שבועות 4–14'),
    ('שלב 3 — בדיקות QA', 'שבועות 15–16'),
    ('שלב 4 — השקה', 'שבוע 17'),
    ('סה״כ לו"ז', '~4 חודשים'),
]
table6 = doc.add_table(rows=1, cols=2)
table6.style = 'Table Grid'
hdr6 = table6.rows[0].cells
for i, h in enumerate(['משך', 'שלב']):
    set_cell_text(hdr6[i], h, bold=True, size=10, bg='2C5594')
    hdr6[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
for idx, (stage, duration) in enumerate(timeline_rows):
    row = table6.add_row().cells
    bg = 'F2F6FF' if idx % 2 == 0 else 'FFFFFF'
    is_total = idx == len(timeline_rows) - 1
    set_cell_text(row[0], duration, bold=is_total, size=10, bg='B8860B' if is_total else bg)
    set_cell_text(row[1], stage, bold=is_total, size=10, bg='B8860B' if is_total else bg)
    if is_total:
        row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

doc.add_paragraph()
add_heading(doc, '5.2 — SLA (הסכם רמת שירות)', 2, ACCENT)
sla_items = [
    'זמן תגובה לבאג קריטי (מערכת לא עולה): עד 4 שעות בימי עסקים.',
    'זמן תגובה לבאג גבוה (פיצ׳ר מרכזי לקוי): עד 24 שעות בימי עסקים.',
    'זמן תגובה לבקשות שיפור: עד 5 ימי עסקים לתכנון.',
    'Uptime מובטח: 99.9% (לא יותר מ-8.7 שעות downtime בשנה).',
    'גיבוי יומי: גיבוי מלא של מסד הנתונים מדי לילה, שמירה ל-30 יום.',
]
for item in sla_items:
    add_bullet(doc, item)

doc.add_paragraph()
add_heading(doc, '5.3 — מה כלול בריטיינר', 2, ACCENT)
included = [
    'תחזוקה שוטפת ועדכוני אבטחה',
    'תיקון באגים שהתגלו לאחר השקה',
    'ניטור ביצועים ואחזקת שרת',
    'חידוש אישורי SSL ודומיין',
    'גיבויים יומיים ובדיקות שחזור רבעוניות',
    'דוחות ביצועים חודשיים',
]
for item in included:
    add_bullet(doc, '✅  ' + item)

add_body(doc, 'לא כלול בריטיינר (דורש הצעת מחיר נפרדת): פיצ׳רים חדשים מהותיים, שינויי עיצוב גדולים, פיתוח אינטגרציות נוספות.', bold=True)

doc.add_paragraph()
add_heading(doc, '5.4 — בעלות ו-IP', 2, ACCENT)
add_body(doc, (
    'עם השלמת התשלום המלא, כל קוד המקור, העיצובים, ונכסי הדיגיטל עוברים לבעלות מלאה של הלקוח. '
    'אנו שומרים את הזכות לציין את הפרויקט ב-Portfolio שלנו (ללא חשיפת נתונים עסקיים רגישים).'
))

add_heading(doc, '5.5 — תנאי ביטול', 2, ACCENT)
cancellation = [
    'ביטול לפני תחילת עבודה: החזר מלא פחות 10% דמי טיפול.',
    'ביטול לאחר סיום שלב 1 (עיצוב): החזר 60% מהתשלום הראשון.',
    'ביטול לאחר סיום שלב 2 (פיתוח): אין החזר על שלבים שהושלמו; הקוד שנכתב מועבר ללקוח.',
]
for item in cancellation:
    add_bullet(doc, item)

add_heading(doc, '5.6 — אחריות', 2, ACCENT)
add_body(doc, 'אחריות לאחר השקה: 60 יום ללא עלות לתיקון כל באג שמקורו בקוד שפיתחנו. לאחר תקופת אחריות זו, הריטיינר החודשי מכסה את כל התחזוקה.')

doc.add_page_break()

# =========================================================
# SECTION 6 — WHY US
# =========================================================
add_heading(doc, 'סעיף 6 — מדוע לבחור בנו?', 1, ACCENT)

why_rows = [
    ('מחיר', 'נמוך', 'גבוה מאוד', 'גבוה — מוצדק'),
    ('גמישות', 'גבוהה', 'נמוכה', 'גבוהה'),
    ('היכרות עם הפרויקט', 'בינוני', 'נמוכה (תחלופה)', 'גבוהה (3 אנשים קבועים)'),
    ('SLA ואחריות', 'לא קיים', 'קיים אך נוקשה', 'ברור ומחויב'),
    ('ניסיון בסאלוני יופי', 'נמוך', 'כללי', 'מותאם לענף'),
    ('בעלות על קוד', 'לפעמים מוגבל', 'לרוב מוגבל', 'בעלות מלאה ללקוח'),
]

table7 = doc.add_table(rows=1, cols=4)
table7.style = 'Table Grid'
hdr7 = table7.rows[0].cells
for i, h in enumerate(['הצוות שלנו', 'סוכנות גדולה', 'פרילנסר זול', 'שיקול']):
    set_cell_text(hdr7[i], h, bold=True, size=10, bg='2C5594')
    hdr7[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

for idx, (label, cheap, agency, us) in enumerate(why_rows):
    row = table7.add_row().cells
    bg = 'F2F6FF' if idx % 2 == 0 else 'FFFFFF'
    set_cell_text(row[0], us, bold=True, size=10, bg='E8F5E9')
    set_cell_text(row[1], agency, size=10, bg=bg)
    set_cell_text(row[2], cheap, size=10, bg=bg)
    set_cell_text(row[3], label, size=10, bg=bg)

doc.add_paragraph()
add_body(doc, '"הבחירה בפרילנסר הזול ביותר היא האסטרטגיה היקרה ביותר לטווח הארוך."', bold=True)

doc.add_page_break()

# =========================================================
# SIGNATURE PAGE
# =========================================================
add_heading(doc, 'אישור וחתימה', 1, ACCENT)
add_body(doc, 'לאישור ההצעה ותחילת ההתקשרות, יש לחתום על הסכם שירותים מלא ולהעביר את התשלום הראשון (40%).')
doc.add_paragraph()

sig_table = doc.add_table(rows=3, cols=2)
sig_table.style = 'Table Grid'
sig_headers = sig_table.rows[0].cells
for i, h in enumerate(['ספק', 'לקוח']):
    set_cell_text(sig_headers[i], h, bold=True, size=11, bg='2C5594')
    sig_headers[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

sig_rows_data = [('שם:', ''), ('חתימה:', ''), ('תאריך:', '')]
for i, (label, _) in enumerate(sig_rows_data):
    row = sig_table.rows[i+1].cells if i < 2 else sig_table.add_row().cells
    set_cell_text(row[0], label + '  __________________', size=11)
    set_cell_text(row[1], label + '  __________________', size=11)

# Fix: add rows properly
sig_table2 = doc.add_table(rows=4, cols=2)
sig_table2.style = 'Table Grid'
headers_sig = sig_table2.rows[0].cells
for i, h in enumerate(['ספק', 'לקוח']):
    set_cell_text(headers_sig[i], h, bold=True, size=11, bg='2C5594')
    headers_sig[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
for i, label in enumerate(['שם', 'חתימה', 'תאריך']):
    row = sig_table2.rows[i+1].cells
    set_cell_text(row[0], f'{label}:  __________________', size=11)
    set_cell_text(row[1], f'{label}:  __________________', size=11)

doc.add_paragraph()
add_body(doc, 'לפרטים נוספים, שאלות, או תיאום פגישת היכרות:', bold=True)
add_body(doc, 'טלפון: ______________    |    מייל: ______________')
doc.add_paragraph()
add_body(doc, 'הצעת מחיר זו הוכנה על בסיס ניתוח שוק מעמיק ומשקפת תמחור ריאלי ועדכני לשנת 2026. תוקף ההצעה: 30 יום ממועד הגשה.')
add_body(doc, 'מסמך זה הוכן על ידי: [שם החברה]   |   © 2026 כל הזכויות שמורות.')

# Save
output_path = r'c:\Users\OWNER\MathU\הצעת_מחיר_סאלון_יופי_2026.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
